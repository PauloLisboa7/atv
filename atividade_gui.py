import tkinter as tk
from tkinter import messagebox
import os
import re

try:
    from docx import Document
except ImportError:
    Document = None


def carregar_atividade1_do_docx(path):
    if Document is None or not os.path.exists(path):
        return None
    try:
        doc = Document(path)
        itens = []
        collecting = False
        current_q = None

        for p in doc.paragraphs:
            t = p.text.strip()
            if not t:
                continue

            if not collecting:
                if t.lower().startswith("após leitura"):
                    collecting = True
                continue

            if t.lower().startswith("parte inferior"):
                break

            t = re.sub(r'^\d+\s+', '', t)

            if re.match(r'^(r|R)\s*=', t):
                resposta = re.sub(r'^(r|R)\s*=\s*', '', t)
                if current_q:
                    itens.append({"pergunta": current_q, "resposta": resposta})
                    current_q = None
                continue

            current_q = t

        return itens if itens else None
    except Exception:
        return None

atividades = {
    1: {
        "titulo": "Atividade 1",
        "itens": [
            {
                "pergunta": "O que você entende por bem-estar pessoal? Você acredita que cuida bem de si mesmo?",
                "resposta": "Bem-estar pessoal é estar bem comigo mesmo, tanto fisicamente quanto mentalmente e emocionalmente. Envolve cuidar da saúde, ter equilíbrio entre estudo, lazer e descanso, além de manter uma boa autoestima."
            },
            {
                "pergunta": "Quais fatores mais influenciam o bem-estar dos jovens atualmente?",
                "resposta": "Os principais fatores são as redes sociais, pressão escolar e familiar, amizades, ambiente em que vivem e questões emocionais como ansiedade e insegurança."
            },
            {
                "pergunta": "Como as redes sociais podem afetar positiva e negativamente a saúde emocional?",
                "resposta": "Positivamente, ajudam na comunicação, aprendizado e entretenimento. Negativamente podem causar comparação excessiva, baixa autoestima, ansiedade e dependência."
            },
            {
                "pergunta": "Você considera suas relações sociais saudáveis? Por quê?",
                "resposta": "Sim, pois procuro manter amizades baseadas no respeito, apoio e confiança, além de evitar conflitos desnecessários."
            },
            {
                "pergunta": "O que pode ser feito para melhorar o bem-estar social entre os jovens?",
                "resposta": "Promover mais diálogo, respeito às diferenças, atividades em grupo, apoio emocional e ambientes acolhedores na escola e na comunidade."
            },
            {
                "pergunta": "Quais atitudes você pode adotar no dia a dia para cuidar melhor da sua saúde física, mental e emocional?",
                "resposta": "Praticar exercícios, manter uma boa alimentação, dormir bem, evitar excesso de redes sociais, organizar a rotina e buscar momentos de lazer e descanso."
            }
        ]
    },
    2: {
        "titulo": "Atividade 2",
        "itens": [
            {"pergunta": "Quanto é 7 x 8?", "resposta": "56"}
        ]
    }
}

doc_path = r"C:\Users\aluno\Downloads\ATV 1 - Paulo Lisboa.docx"

dados_docx = carregar_atividade1_do_docx(doc_path)
if dados_docx:
    atividades[1]["itens"] = dados_docx


class App:
    def __init__(self, root):
        self.root = root
        self.root.title("Minhas Atividades")
        self.root.geometry("500x400")

        self.main_frame = tk.Frame(root)
        self.main_frame.pack(expand=True, fill="both")

        self.start_button = tk.Button(
            self.main_frame,
            text="CLIQUE AQUI",
            font=("Arial", 24, "bold"),
            bg="#4CAF50",
            fg="white",
            padx=20,
            pady=10,
            command=self.show_atividades
        )
        self.start_button.place(relx=0.5, rely=0.5, anchor="center")

        self.activity_label = None
        self.answer_text = None
        self.export_button = None

    def show_atividades(self):
        self.start_button.destroy()

        if self.activity_label:
            self.activity_label.destroy()

        self.activity_label = tk.Label(
            self.main_frame,
            text="Escolha uma atividade:",
            font=("Arial", 18)
        )
        self.activity_label.pack(pady=10)

        btn_frame = tk.Frame(self.main_frame)
        btn_frame.pack(pady=10)

        for key in sorted(atividades.keys()):
            bt = tk.Button(
                btn_frame,
                text=f"{atividades[key]['titulo']}",
                font=("Arial", 14),
                width=15,
                command=lambda k=key: self.mostrar_resposta(k)
            )
            bt.pack(side="left", padx=10)

    def mostrar_resposta(self, key):
        if self.answer_text:
            self.answer_text.destroy()
        if self.export_button:
            self.export_button.destroy()

        self.answer_text = tk.Text(self.main_frame, font=("Arial", 13), wrap="word", height=14, width=58)
        self.answer_text.pack(pady=10)
        self.answer_text.configure(state="normal")

        for idx, item in enumerate(atividades[key]["itens"], start=1):
            self.answer_text.insert("end", f"{idx}. Pergunta: {item['pergunta']}\n")
            self.answer_text.insert("end", f"   Resposta: {item['resposta']}\n\n")

        self.answer_text.configure(state="disabled")

        self.export_button = tk.Button(
            self.main_frame,
            text="Exportar para Word",
            font=("Arial", 14, "bold"),
            bg="#2196F3",
            fg="white",
            command=lambda k=key: self.exportar_para_word(k)
        )
        self.export_button.pack(pady=10)

    def exportar_para_word(self, key):
        if Document is None:
            messagebox.showerror(
                "Erro",
                "Biblioteca python-docx não está instalada.\nInstale com: pip install python-docx"
            )
            return

        doc = Document()
        doc.add_heading(atividades[key]['titulo'], level=1)

        for idx, item in enumerate(atividades[key]["itens"], start=1):
            doc.add_paragraph(f"{idx}. Pergunta: {item['pergunta']}")
            doc.add_paragraph(f"Resposta: {item['resposta']}")

        filename = os.path.join(os.path.expanduser("~"), f"atividade_{key}.docx")
        try:
            doc.save(filename)
            messagebox.showinfo("Sucesso", f"Arquivo salvo em: {filename}")
            os.startfile(filename)
        except Exception as e:
            messagebox.showerror("Erro", f"Falha ao salvar o arquivo: {e}")


if __name__ == "__main__":
    root = tk.Tk()
    app = App(root)
    root.mainloop()
